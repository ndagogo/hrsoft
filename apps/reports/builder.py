"""Execute custom reports and export to CSV, Excel, PDF, Word."""

import io
from datetime import date, datetime
from decimal import Decimal

from django.http import HttpResponse

from apps.core.exports import build_csv_response, build_excel_response


def _serialize_value(val):
    if val is None:
        return ""
    if isinstance(val, (date, datetime)):
        return val.isoformat()
    if isinstance(val, Decimal):
        return str(val)
    return str(val)


def run_report(report_def):
    from apps.employees.models import Employee, EmploymentStatus
    from apps.attendance.models import AttendanceRecord
    from apps.leave.models import LeaveRequest
    from apps.payroll.models import Payslip

    rtype = report_def.report_type
    filters = report_def.filters or {}

    if rtype == "employees":
        qs = Employee.objects.select_related("user", "department", "designation")
        if filters.get("status"):
            qs = qs.filter(status=filters["status"])
        if filters.get("department"):
            qs = qs.filter(department_id=filters["department"])
        default_cols = ["employee_id", "full_name", "department", "designation", "status", "basic_salary"]
        cols = report_def.columns or default_cols
        rows = []
        for emp in qs:
            row = []
            for col in cols:
                if col == "full_name":
                    row.append(emp.full_name)
                elif col == "department":
                    row.append(emp.department.name if emp.department else "")
                elif col == "designation":
                    row.append(emp.designation.title if emp.designation else "")
                else:
                    row.append(_serialize_value(getattr(emp, col, "")))
            rows.append(row)
        return cols, rows

    if rtype == "attendance":
        qs = AttendanceRecord.objects.select_related("employee__user", "employee__department")
        if filters.get("date_from"):
            qs = qs.filter(date__gte=filters["date_from"])
        if filters.get("date_to"):
            qs = qs.filter(date__lte=filters["date_to"])
        cols = report_def.columns or ["employee_id", "employee_name", "date", "status", "worked_hours"]
        rows = []
        for rec in qs[:5000]:
            rows.append([
                rec.employee.employee_id,
                rec.employee.full_name,
                rec.date.isoformat(),
                rec.get_status_display(),
                str(rec.worked_hours),
            ])
        return cols, rows

    if rtype == "leave":
        qs = LeaveRequest.objects.select_related("employee__user", "leave_type")
        if filters.get("status"):
            qs = qs.filter(status=filters["status"])
        cols = report_def.columns or ["employee", "leave_type", "start_date", "end_date", "status"]
        rows = [[r.employee.full_name, r.leave_type.name, r.start_date.isoformat(), r.end_date.isoformat(), r.get_status_display()] for r in qs[:2000]]
        return cols, rows

    if rtype == "payroll":
        qs = Payslip.objects.select_related("employee__user", "period")
        if filters.get("period"):
            qs = qs.filter(period_id=filters["period"])
        cols = report_def.columns or ["employee_id", "period", "gross_pay", "net_pay"]
        rows = [[s.employee.employee_id, s.period.name, str(s.gross_pay), str(s.net_pay)] for s in qs[:2000]]
        return cols, rows

    return ["info"], [["No data for this report type"]]


def export_report_pdf(report_def, headers, rows):
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
    from reportlab.lib.styles import getSampleStyleSheet

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=landscape(A4))
    styles = getSampleStyleSheet()
    elements = [Paragraph(report_def.name, styles["Heading1"])]
    data = [headers] + rows[:500]
    table = Table(data, repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#111a2e")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
    ]))
    elements.append(table)
    doc.build(elements)
    buffer.seek(0)
    response = HttpResponse(buffer.getvalue(), content_type="application/pdf")
    response["Content-Disposition"] = f'attachment; filename="{report_def.name.replace(" ", "_")}.pdf"'
    return response


def export_report_docx(report_def, headers, rows):
    from docx import Document

    doc = Document()
    doc.add_heading(report_def.name, 0)
    if report_def.description:
        doc.add_paragraph(report_def.description)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = str(h)
    for row in rows[:500]:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    response = HttpResponse(
        buffer.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    response["Content-Disposition"] = f'attachment; filename="{report_def.name.replace(" ", "_")}.docx"'
    return response


def export_report(report_def, fmt):
    headers, rows = run_report(report_def)
    if fmt == "pdf":
        return export_report_pdf(report_def, headers, rows)
    if fmt == "docx":
        return export_report_docx(report_def, headers, rows)
    if fmt == "xlsx":
        return build_excel_response(f"{report_def.name}.xlsx", headers, rows)
    return build_csv_response(f"{report_def.name}.csv", headers, rows)
